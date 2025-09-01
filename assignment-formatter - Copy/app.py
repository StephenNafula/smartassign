import os
from functools import wraps
import json # Import json for parsing features
from flask import Flask, render_template, request, redirect, url_for, session, flash, g, jsonify, send_file
from flask_bcrypt import Bcrypt
import mysql.connector
import html # optional for sanitization

from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Pt, RGBColor
from bs4 import BeautifulSoup
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from datetime import datetime, timedelta
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

# Configure WeasyPrint logging for more verbosity
logging.getLogger('weasyprint').setLevel(logging.DEBUG)

# -------- Config ----------
app = Flask(__name__, template_folder="templates", static_folder="static")
app.secret_key = os.getenv("FLASK_SECRET_KEY", "dev-secret-change-me")
bcrypt = Bcrypt(app)

# WeasyPrint DLL directories for Windows (if not set as system env variable)
# Example: WEASYPRINT_DLL_DIRECTORIES = ["C:\msys64\mingw64\bin"]
WEASYPRINT_DLL_DIRECTORIES = os.getenv("WEASYPRINT_DLL_DIRECTORIES", "").split(";")
if WEASYPRINT_DLL_DIRECTORIES == [""]:
    WEASYPRINT_DLL_DIRECTORIES = []
os.environ["PATH"] += os.pathsep + os.pathsep.join(WEASYPRINT_DLL_DIRECTORIES)

# Placeholder for payment gateway API keys (e.g., Flutterwave, IntaSend)
API_KEYS = {
    "FLUTTERWAVE_PUBLIC": os.getenv("FLUTTERWAVE_PUBLIC_KEY"),
    "FLUTTERWAVE_SECRET": os.getenv("FLUTTERWAVE_SECRET_KEY"),
    "INTASEND_PUBLIC": os.getenv("INTASEND_PUBLIC_KEY"),
    "INTASEND_SECRET": os.getenv("INTASEND_SECRET_KEY"),
}

DB_CONFIG = {
    "host": os.getenv("DB_HOST", "localhost"),
    "user": os.getenv("DB_USER", "formatter_user"),
    "password": os.getenv("DB_PASS", "strongpassword123"),
    "database": os.getenv("DB_NAME", "assignment_formatter_db"),
    "autocommit": True,
    "ssl_disabled": True # Disable SSL for local development
}

def get_db_conn():
    # simple single connection for dev; for production use connection pool
    if not hasattr(g, "db_conn"):
        try:
            g.db_conn = mysql.connector.connect(**DB_CONFIG)
        except mysql.connector.errors.InterfaceError as ie:
            # Handle MySQL 8+ caching_sha2_password policy which may require SSL
            # If server requires secure connection, try a fallback that enables SSL
            # and allows public key retrieval (useful for local/dev setups).
            try:
                if getattr(ie, 'errno', None) == 2061 or 'caching_sha2_password' in str(ie):
                    logging.warning("DB connection failed with InterfaceError; attempting SSL-enabled fallback: %s", ie)
                    # First fallback: enable SSL but don't verify cert (for local dev only)
                    fallback = DB_CONFIG.copy()
                    fallback.update({
                        'ssl_disabled': False,
                        'ssl_verify_cert': False,
                    })
                    try:
                        g.db_conn = mysql.connector.connect(**fallback)
                    except Exception:
                        logging.warning("SSL-enabled fallback failed; attempting auth_plugin=mysql_native_password fallback")
                        # Second fallback: try using mysql_native_password auth plugin (local/dev)
                        fb2 = DB_CONFIG.copy()
                        fb2.update({'auth_plugin': 'mysql_native_password'})
                        g.db_conn = mysql.connector.connect(**fb2)
                else:
                    raise
            except Exception:
                # Re-raise the original interface error for visibility
                logging.exception("DB connection fallback also failed")
                raise
    return g.db_conn

def get_db_cursor():
    conn = get_db_conn()
    return conn.cursor(dictionary=True)

@app.teardown_appcontext
def close_db_conn(exc):
    if hasattr(g, "db_conn"):
        try:
            g.db_conn.close()
        except:
            pass

# --------- Helpers ----------
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login', next=request.path))
        return f(*args, **kwargs)
    return decorated

# Simple heading detection (Stage 1) - Re-added
def detect_sections(text):
    lines = [l.strip() for l in text.splitlines() if l.strip() != ""]
    sections = []
    cur_heading = "Body"
    cur_body = []
    keywords = {'introduction','abstract','conclusion','methodology','results','references'}
    for line in lines:
        low = line.lower()
        if len(line.split()) <= 6 and (low in keywords or low.endswith(':') or line.isupper()):
            if cur_body:
                sections.append((cur_heading, "\n".join(cur_body)))
                cur_body = []
            cur_heading = line
        else:
            cur_body.append(line)
    if cur_body:
        sections.append((cur_heading, "\n".join(cur_body)))
    return sections

# Helper to get the user's current plan details
def get_user_current_plan(user_id):
    cur = get_db_cursor()
    cur.execute("""
        SELECT u.id as user_id, u.name, u.email, u.document_credits,
               p.id as plan_id, p.name as plan_name, p.plan_type, p.price, p.currency, p.features_json, p.is_watermarked_export, p.document_cost, p.ai_features_enabled, p.max_documents_per_day
        FROM users u
        JOIN plans p ON u.current_plan_id = p.id
        WHERE u.id = %s
    """, (user_id,))
    user_plan = cur.fetchone()
    cur.close()
    if user_plan and user_plan['features_json']:
        user_plan['features'] = json.loads(user_plan['features_json'])
    else:
        user_plan['features'] = []
    return user_plan

# --------- Routes ----------
@app.route("/")
def index():
    # If logged in, redirect to dashboard
    if 'user_id' in session:
        return redirect(url_for('dashboard'))
    return render_template("landing.html")

# Registration
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        name = request.form.get("name", "").strip()
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")

        if not name or not email or not password:
            flash("All fields required", "error")
            return redirect(url_for("register"))

        conn = get_db_conn()
        cur = conn.cursor(dictionary=True)

        # check existing
        cur.execute("SELECT id FROM users WHERE email=%s", (email,))
        existing = cur.fetchone()
        if existing:
            flash("Email already registered. Please log in.", "error")
            cur.close()
            return redirect(url_for("login"))

        pw_hash = bcrypt.generate_password_hash(password).decode("utf-8")
        cur.execute("INSERT INTO users (name, email, password_hash) VALUES (%s,%s,%s)",
                    (name, email, pw_hash))
        conn.commit()
        user_id = cur.lastrowid
        cur.close()

        # log in user
        session['user_id'] = user_id
        session['user_name'] = name
        flash("Registration successful. Welcome!", "success")
        return redirect(url_for("dashboard"))

    return render_template("register.html")

# Login
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")

        conn = get_db_conn()
        cur = conn.cursor(dictionary=True)
        cur.execute("SELECT id, name, password_hash FROM users WHERE email=%s", (email,))
        user = cur.fetchone()
        cur.close()

        if not user:
            flash("Invalid credentials", "error")
            return redirect(url_for("login"))

        if not bcrypt.check_password_hash(user["password_hash"], password):
            flash("Invalid credentials", "error")
            return redirect(url_for("login"))

        session['user_id'] = user['id']
        session['user_name'] = user['name']
        flash("Signed in successfully", "success")
        next_url = request.args.get('next') or url_for('dashboard')
        return redirect(next_url)

    return render_template("login.html")

# Logout
@app.route("/logout", methods=["POST"])
def logout():
    session.pop('user_id', None)
    session.pop('user_name', None) # Clear user_name as well
    flash("Logged out", "info")
    return redirect(url_for("login"))

# Simple dashboard (requires login)
@app.route("/dashboard")
@login_required
def dashboard():
    user_id = session.get('user_id')
    if not user_id:
        flash("Please log in to view your dashboard.", "error")
        return redirect(url_for('login'))

    # Get user's current plan and credits
    user_plan_info = get_user_current_plan(user_id)

    cur = get_db_cursor()
    cur.execute("SELECT id, title, status, created_at, updated_at FROM documents WHERE user_id=%s ORDER BY updated_at DESC", (user_id,))
    documents = cur.fetchall()
    cur.close()
    return render_template("dashboard.html", user_name=session.get("user_name"), documents=documents, user_plan_info=user_plan_info)

@app.route('/editor')
@app.route('/editor/<int:doc_id>')
def editor(doc_id=None):
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user_id = session.get('user_id')
    user_plan_info = get_user_current_plan(user_id)

    return render_template('editor.html', doc_id=doc_id, user_plan_info=user_plan_info)

@app.route('/save_assignment', methods=['POST'])
def save_assignment():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    content = request.form.get("content")  # The HTML content from Quill.js
    session['assignment_content'] = content  # Save the HTML content
    flash("✅ Assignment saved successfully!", "success")
    return redirect(url_for("dashboard"))

# Stage 1 Preview Endpoint - Re-added
@app.route("/api/doc/preview", methods=["POST"])
def preview():
    data = request.json or {}
    text = data.get("text", "")
    detect = data.get("detect", True)
    sections = detect_sections(text) if detect else [("Body", text)]
    html = "<!doctype html><html><head><meta charset='utf-8'><title>Preview</title>"
    html += "<style>body{font-family: 'Times New Roman', serif; font-size:12pt; margin:0.7in} h2{font-weight:bold}</style>"
    html += "</head><body>"
    for h, b in sections:
        html += f"<h2>{h}</h2>"
        for paragraph in b.split("\n\n"):
            html += f"<p>{paragraph.replace(chr(10), '<br>')}</p>"
    html += "</body></html>"
    return jsonify({"html": html})

@app.route("/health")
def health():
    return "ok", 200

@app.route("/api/save_draft", methods=["POST"])
def api_save_draft():
    if 'user_id' not in session:
        return jsonify({"error": "unauthenticated"}), 401
    data = request.json or {}
    content = data.get('content', '')
    title = data.get('title', '')[:250]  # max length safeguard
    user_id = session['user_id']

    cur = get_db_cursor()
    # If client has document_id, update; else insert new draft
    doc_id = data.get('document_id')
    if doc_id:
        cur.execute("UPDATE documents SET content=%s, title=%s WHERE id=%s AND user_id=%s",
                    (content, title, doc_id, user_id))
        get_db_conn().commit()
        return jsonify({"ok": True, "document_id": doc_id})
    else:
        cur.execute("INSERT INTO documents (user_id, title, content, status) VALUES (%s,%s,%s,%s)",
                    (user_id, title, content, 'draft'))
        get_db_conn().commit()
        new_id = cur.lastrowid
        return jsonify({"ok": True, "document_id": new_id})

@app.route("/api/submit_assignment", methods=["POST"])
def api_submit_assignment():
    if 'user_id' not in session:
        return jsonify({"error": "unauthenticated"}), 401
    data = request.json or {}
    document_id = data.get("document_id")
    if not document_id:
        return jsonify({"error": "missing document_id"}), 400

    cur = get_db_cursor()
    # mark as submitted
    cur.execute("UPDATE documents SET status=%s WHERE id=%s AND user_id=%s",
                ('submitted', document_id, session['user_id']))
    get_db_conn().commit()
    return jsonify({"ok": True, "document_id": document_id})

@app.route("/api/doc/<int:doc_id>")
def api_get_document(doc_id):
    if 'user_id' not in session:
        return jsonify({"error": "unauthenticated"}), 401

    cur = get_db_cursor()
    cur.execute("SELECT id, title, content, status, created_at FROM documents WHERE id=%s AND user_id=%s",
                (doc_id, session['user_id']))
    document = cur.fetchone()

    if not document:
        return jsonify({"ok": False, "error": "Document not found or unauthorized"}), 404

    return jsonify({"ok": True, "document": document})

@app.route('/export/<int:doc_id>')
def export_document(doc_id):
    if 'user_id' not in session:
        return "Unauthorized", 401

    user_id = session.get('user_id')
    user_plan_info = get_user_current_plan(user_id)

    if not user_plan_info:
        return "User plan not found.", 400

    # Check for document limits (for free and standard plans)
    if user_plan_info['plan_type'] == 'free' and user_plan_info['document_credits'] <= 0:
        flash("You have no free document credits left for today. Please upgrade your plan or wait until tomorrow.", "error")
        print("DEBUG: Free plan, no credits left.") # Debug print
        return redirect(url_for('editor', doc_id=doc_id))
    elif user_plan_info['plan_type'] == 'one_time_document' and user_plan_info['document_credits'] <= 0:
        flash("You have no document credits. Please purchase more to export.", "error")
        print("DEBUG: One-time plan, no credits left.") # Debug print
        return redirect(url_for('editor', doc_id=doc_id))
    elif user_plan_info['plan_type'] == 'monthly_subscription' and user_plan_info['max_documents_per_day'] is not None:
        # For simplicity, we'll implement a basic check here.
        # A more robust solution would involve a dedicated 'document_exports' table or a more complex credit system.
        # For now, let's assume a daily count for 'max_documents_per_day'
        cur = get_db_cursor()
        today_start = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        cur.execute("SELECT COUNT(*) as count FROM documents WHERE user_id=%s AND updated_at >= %s AND status = 'exported'", (user_id, today_start))
        exported_today = cur.fetchone()['count']
        cur.close()
        if exported_today >= user_plan_info['max_documents_per_day']:
            flash(f"You have reached your daily export limit of {user_plan_info['max_documents_per_day']} documents. Please upgrade your plan or try again tomorrow.", "error")
            print(f"DEBUG: Monthly plan, daily limit reached: {exported_today}") # Debug print
            return redirect(url_for('editor', doc_id=doc_id))


    typ = request.args.get('type', 'pdf').lower()
    print(f"DEBUG: Export type requested: {typ}") # Debug print
    cur = get_db_cursor()
    cur.execute("SELECT * FROM documents WHERE id=%s AND user_id=%s", (doc_id, session.get('user_id')))
    doc = cur.fetchone()
    if not doc:
        print(f"DEBUG: Document {doc_id} not found or unauthorized for user {session.get('user_id')}") # Debug print
        return "Not found or unauthorized", 404

    html_content = doc['content'] or ''
    print(f"DEBUG: Document content length: {len(html_content)}") # Debug print
    
    watermark_style = ""
    if user_plan_info['is_watermarked_export']:
        watermark_style = """
          body::after {
            content: "Assignment Formatter - Watermark";
            position: fixed;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 3em;
            color: rgba(0, 0, 0, 0.1);
            z-index: 9999;
            pointer-events: none;
            white-space: nowrap;
          }
        """
        print(f"DEBUG: Watermark enabled: {user_plan_info['is_watermarked_export']}") # Debug print

    # Build a minimal HTML wrapper matching preview page CSS for PDF reproducibility
    full_html = f"""
    <html>
      <head>
        <meta charset="utf-8"/>
        <style>
          body{{ font-family: "Times New Roman", serif; font-size:12pt; margin:1in; line-height:1.5; color:#111; }}
          h1,h2,h3{{ font-weight:bold; }}
          p{{ margin: 0 0 0.8em 0; }}
          {watermark_style}
        </style>
      </head>
      <body>{html_content}</body>
    </html>
    """
    print("DEBUG: Full HTML prepared for export.") # Debug print

    try:
        if typ == 'pdf':
            print("DEBUG: Attempting PDF generation.") # Debug print
            try:
                pdf_bytes = HTML(string=full_html).write_pdf(stylesheets=[CSS(string='@page { size: A4; margin: 1in }')])
                print("DEBUG: PDF generated by WeasyPrint.")
            except Exception as wp_err:
                # Log WeasyPrint error and attempt a simple ReportLab fallback
                print(f"WARNING: WeasyPrint PDF generation failed: {wp_err}")
                print("DEBUG: Falling back to ReportLab for a simple PDF.")
                buf = io.BytesIO()
                c = canvas.Canvas(buf, pagesize=A4)
                width, height = A4
                # Basic rendering: write plain text paragraphs
                y = height - inch
                text_obj = c.beginText(inch, y)
                text_obj.setFont('Times-Roman', 12)
                for line in BeautifulSoup(html_content, 'html.parser').get_text(separator='\n').splitlines():
                    text_obj.textLine(line)
                c.drawText(text_obj)
                c.showPage()
                c.save()
                buf.seek(0)
                pdf_bytes = buf.read()
                print("DEBUG: PDF generated by ReportLab fallback.")
            
            # Decrement credits or update daily count AFTER successful export
            if user_plan_info['plan_type'] == 'one_time_document':
                cur = get_db_cursor()
                cur.execute("UPDATE users SET document_credits = document_credits - 1 WHERE id = %s", (user_id,))
                get_db_conn().commit()
                cur.close()
                print("DEBUG: Decremented one-time document credit.") # Debug print
            elif user_plan_info['plan_type'] == 'free' and user_plan_info['document_credits'] > 0:
                 cur = get_db_cursor()
                 cur.execute("UPDATE users SET document_credits = document_credits - 1 WHERE id = %s", (user_id,))
                 get_db_conn().commit()
                 cur.close()
                 print("DEBUG: Decremented free document credit.") # Debug print

            # Mark document as exported if not already
            cur = get_db_cursor()
            cur.execute("UPDATE documents SET status = %s WHERE id = %s", ('exported', doc_id))
            get_db_conn().commit()
            cur.close()
            print("DEBUG: Document status updated to exported.") # Debug print

            return send_file(io.BytesIO(pdf_bytes), as_attachment=True, download_name=f"assignment_{doc_id}.pdf", mimetype='application/pdf')

        elif typ == 'docx':
            print("DEBUG: Attempting DOCX generation.") # Debug print
            # Parse HTML and construct .docx
            soup = BeautifulSoup(html_content, 'html.parser')
            docx = Document()
            # Set default style font (python-docx has limitations on full style control, but set Normal)
            style = docx.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)

            # Apply watermark if needed (python-docx doesn't directly support background watermarks like CSS)
            # For DOCX, a watermark would typically be added as a header/footer image or a Shape.
            # This is a simplification. A real implementation would involve more complex docx manipulation.
            if user_plan_info['is_watermarked_export']:
                # Placeholder for DOCX watermark: add a paragraph with light grey text
                header = docx.sections[0].header
                paragraph = header.paragraphs[0]
                run = paragraph.add_run("Assignment Formatter - Watermark")
                run.font.color.rgb = RGBColor(192, 192, 192) # Light grey
                run.font.size = Pt(24)
                print("DEBUG: DOCX watermark applied.") # Debug print

            # Convert headings and paragraphs
            for elem in soup.find_all(['h1','h2','h3','p','ul','ol']):
                if elem.name in ['h1','h2','h3']:
                    level = 1 if elem.name=='h1' else (2 if elem.name=='h2' else 3)
                    docx.add_heading(elem.get_text(), level=level)
                elif elem.name == 'p':
                    p = docx.add_paragraph(elem.get_text())
                elif elem.name in ['ul','ol']:
                    for li in elem.find_all('li'):
                        p = docx.add_paragraph(li.get_text(), style='List Bullet' if elem.name=='ul' else 'List Number')

            f = io.BytesIO()
            docx.save(f)
            f.seek(0)
            print("DEBUG: DOCX generated.") # Debug print

            # Decrement credits or update daily count AFTER successful export
            if user_plan_info['plan_type'] == 'one_time_document':
                cur = get_db_cursor()
                cur.execute("UPDATE users SET document_credits = document_credits - 1 WHERE id = %s", (user_id,))
                get_db_conn().commit()
                cur.close()
                print("DEBUG: Decremented one-time document credit.") # Debug print
            elif user_plan_info['plan_type'] == 'free' and user_plan_info['document_credits'] > 0:
                 cur = get_db_cursor()
                 cur.execute("UPDATE users SET document_credits = document_credits - 1 WHERE id = %s", (user_id,))
                 get_db_conn().commit()
                 cur.close()
                 print("DEBUG: Decremented free document credit.") # Debug print

            # Mark document as exported if not already
            cur = get_db_cursor()
            cur.execute("UPDATE documents SET status = %s WHERE id = %s", ('exported', doc_id))
            get_db_conn().commit()
            cur.close()
            print("DEBUG: Document status updated to exported.") # Debug print

            return send_file(f, as_attachment=True, download_name=f"assignment_{doc_id}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        else:
            print(f"DEBUG: Unsupported export type: {typ}") # Debug print
            return "Unsupported export type", 400
    except Exception as e:
        get_db_conn().rollback() # Rollback in case of error
        print(f"ERROR: Exception during export: {e}") # Debug print
        return jsonify({"ok": False, "error": str(e)}), 500

@app.route("/plans")
def plans():
    cur = get_db_cursor()
    cur.execute("SELECT id, name, plan_type, price, currency, features_json, is_watermarked_export, document_cost, ai_features_enabled, max_documents_per_day, initial_credits FROM plans ORDER BY id ASC")
    plans_data = cur.fetchall()
    cur.close()

    # Parse features_json string into a Python list for each plan
    for plan in plans_data:
        if plan['features_json']:
            plan['features'] = json.loads(plan['features_json'])
        else:
            plan['features'] = []

    return render_template("plans.html", plans=plans_data)

@app.route("/checkout")
def checkout():
    plan_id = request.args.get('plan_id', type=int)
    selected_plan = None

    if plan_id:
        cur = get_db_cursor()
        cur.execute("SELECT id, name, plan_type, price, currency, features_json, is_watermarked_export, document_cost, ai_features_enabled, max_documents_per_day, initial_credits FROM plans WHERE id = %s", (plan_id,))
        selected_plan = cur.fetchone()
        cur.close()

        if selected_plan and selected_plan['features_json']:
            selected_plan['features'] = json.loads(selected_plan['features_json'])
        else:
            selected_plan['features'] = []

    return render_template("checkout.html", selected_plan=selected_plan)


# Silent handler for noisy external tracker requests that hit /hybridaction/*
# Returns 204 No Content so the logs are not filled with 404s during development.
@app.route('/hybridaction/<path:_subpath>', methods=['GET', 'POST', 'OPTIONS'])
def _hybridaction_handler(_subpath):
    return ('', 204)

# API for initiating payment
@app.route("/api/initiate_payment", methods=["POST"])
@login_required
def initiate_payment():
    user_id = session.get('user_id')
    if not user_id: # Should be caught by login_required, but for safety
        return jsonify({"error": "unauthenticated"}), 401

    data = request.json or {}
    plan_id = data.get('plan_id', type=int)
    quantity = data.get('quantity', 1) # For pay-per-document

    if not plan_id:
        return jsonify({"error": "Plan ID is required"}), 400

    cur = get_db_cursor()
    cur.execute("SELECT id, name, plan_type, price, currency, document_cost FROM plans WHERE id = %s", (plan_id,))
    plan = cur.fetchone()
    cur.close()

    if not plan:
        return jsonify({"error": "Plan not found"}), 404

    amount = 0
    if plan['plan_type'] == 'one_time_document':
        amount = plan['document_cost'] * quantity
    elif plan['plan_type'] == 'monthly_subscription':
        amount = plan['price']
    else:
        return jsonify({"error": "Invalid plan type for payment"}), 400

    # For now, simulate a successful payment and create a transaction record
    try:
        conn = get_db_conn()
        # Create a pending transaction
        cur = conn.cursor(dictionary=True)
        cur.execute("INSERT INTO transactions (user_id, plan_id, amount, currency, status) VALUES (%s, %s, %s, %s, %s)",
                    (user_id, plan['id'], amount, plan['currency'], 'pending'))
        transaction_id = cur.lastrowid
        conn.commit()
        # Simulate successful payment immediately for development
        # In a real app, this would be a callback from the payment gateway
        # For now, we'll directly call the logic to update user plan/credits
        # You will replace this with actual payment gateway integration.

        # Update user's plan/credits
        if plan['plan_type'] == 'one_time_document':
            cur.execute("UPDATE users SET document_credits = document_credits + %s WHERE id = %s",
                        (quantity, user_id))
        elif plan['plan_type'] == 'monthly_subscription':
            # Set subscription to end one month from now
            # For simplicity, not handling existing subscriptions here, just overwriting
            cur.execute("UPDATE users SET current_plan_id = %s, subscription_end_date = DATE_ADD(NOW(), INTERVAL 1 MONTH) WHERE id = %s",
                        (plan['id'], user_id))
        
        # Update transaction status to completed
        cur.execute("UPDATE transactions SET status = %s, payment_gateway_ref = %s WHERE id = %s",
                    ('completed', f"SIMULATED_REF_{transaction_id}", transaction_id))
        conn.commit()
        cur.close()

        return jsonify({"ok": True, "message": "Payment simulated successfully and plan updated!", "transaction_id": transaction_id}), 200

    except Exception as e:
        get_db_conn().rollback() # Rollback in case of error
        return jsonify({"ok": False, "error": str(e)}), 500

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0")
